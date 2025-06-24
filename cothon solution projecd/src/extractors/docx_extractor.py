"""
DOCX document extractor module for legal contract analysis.
Handles text extraction and basic structure analysis from DOCX files.
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile
import io
from datetime import datetime
import json

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extracts text and structure from DOCX documents."""

    # Map WD_PARAGRAPH_ALIGNMENT to human-readable strings
    ALIGNMENT_MAP = {
        WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
        WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
        WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
        WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
        None: "default"
    }

    def __init__(self):
        """Initialize the DOCX extractor."""
        self.section_patterns = [
            r'^(\d+\.\s*[A-Z][^\n]+)',  # Numbered sections (e.g., "1. Introduction")
            r'^([A-Z][A-Z\s]+:)',  # ALL CAPS sections (e.g., "ARTICLE I:")
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*:)',  # Title Case sections (e.g., "Payment Terms:")
        ]

    def extract_text(self, file_path: str) -> Dict[str, any]:
        """
        Extract text and metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary containing:
            - text: Full document text
            - metadata: Document metadata
            - sections: Identified sections
            - paragraphs: List of paragraphs with formatting
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            if file_path.suffix.lower() != '.docx':
                raise ValueError(f"Invalid file type: {file_path}. Expected .docx")

            # Load document
            doc = Document(file_path)

            # Extract metadata
            metadata = self._extract_metadata(doc)

            # Extract paragraphs with formatting
            paragraphs = []
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    # Get paragraph formatting
                    para_info = {
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                        "alignment": self.ALIGNMENT_MAP.get(para.alignment, "default"),
                        "is_heading": para.style.name.startswith("Heading") if para.style else False,
                        "runs": []
                    }

                    # Get run formatting
                    for run in para.runs:
                        run_info = {
                            "text": run.text,
                            "bold": bool(run.bold),
                            "italic": bool(run.italic),
                            "underline": bool(run.underline),
                            "font_size": run.font.size.pt if run.font.size else None,
                            "font_name": run.font.name if run.font.name else None
                        }
                        para_info["runs"].append(run_info)

                    paragraphs.append(para_info)
                    full_text.append(para.text)

            # Combine all text
            text = "\n".join(full_text)

            # Identify sections
            sections = self._identify_sections(text, paragraphs)

            return {
                "text": text,
                "metadata": metadata,
                "sections": sections,
                "paragraphs": paragraphs,
                "total_paragraphs": len(paragraphs)
            }

        except FileNotFoundError as e:
            logger.error(f"File not found: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise

    def _extract_metadata(self, doc: Document) -> Dict[str, any]:
        """
        Extract metadata from DOCX document.

        Args:
            doc: python-docx Document object

        Returns:
            Dictionary of metadata
        """
        metadata = {}

        try:
            # Core properties
            core_props = doc.core_properties
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "created": core_props.created.isoformat() if core_props.created else "",
                "modified": core_props.modified.isoformat() if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
                "category": core_props.category or "",
                "keywords": core_props.keywords or "",
                "subject": core_props.subject or ""
            })

            # Document statistics
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "table_count": len(doc.tables)
            })

        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}")

        return metadata

    def _identify_sections(self, text: str, paragraphs: List[Dict]) -> List[Dict[str, any]]:
        """
        Identify document sections using pattern matching and style analysis.

        Args:
            text: Document text
            paragraphs: List of paragraph information

        Returns:
            List of identified sections with their text
        """
        sections = []
        current_section = None
        current_text = []

        # First, try to identify sections using style-based headings
        for para in paragraphs:
            if para["is_heading"]:
                if current_section:
                    sections.append({
                        "title": current_section,
                        "text": "\n".join(current_text).strip()
                    })
                current_section = para["text"].strip()
                current_text = [para["text"]]
            elif current_section:
                current_text.append(para["text"])

        # If no style-based headings found, fall back to pattern matching
        if not sections:
            lines = text.split("\n")
            for line in lines:
                is_section = False
                for pattern in self.section_patterns:
                    match = re.match(pattern, line.strip())
                    if match:
                        if current_section:
                            sections.append({
                                "title": current_section,
                                "text": "\n".join(current_text).strip()
                            })
                        current_section = match.group(1).strip()
                        current_text = [line]
                        is_section = True
                        break
                if not is_section and current_section:
                    current_text.append(line)

        # Add last section
        if current_section:
            sections.append({
                "title": current_section,
                "text": "\n".join(current_text).strip()
            })

        return sections

    def extract_tables(self, file_path: str) -> List[Dict[str, any]]:
        """
        Extract tables from DOCX document.

        Args:
            file_path: Path to the DOCX file

        Returns:
            List of extracted tables
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            doc = Document(file_path)
            tables = []

            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data and any(cell for row in table_data for cell in row):  # Skip empty tables
                    tables.append({
                        "table_number": table_num,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0,
                        "data": table_data
                    })

            return tables

        except FileNotFoundError as e:
            logger.error(f"File not found: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error extracting tables from DOCX {file_path}: {str(e)}")
            raise

    def extract_images(self, file_path: str, output_dir: Optional[str] = None) -> List[Dict[str, any]]:
        """
        Extract images from DOCX document.

        Args:
            file_path: Path to the DOCX file
            output_dir: Directory to save extracted images

        Returns:
            List of extracted image information
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            images = []
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # List all files in the DOCX archive
                for file_info in docx_zip.infolist():
                    if file_info.filename.startswith('word/media/'):
                        # Extract image data
                        image_data = docx_zip.read(file_info.filename)
                        image_ext = Path(file_info.filename).suffix[1:].lower()
                        if image_ext in ('png', 'jpg', 'jpeg', 'gif'):
                            image_info = {
                                "image_number": len(images) + 1,
                                "format": image_ext
                            }
                            if output_dir:
                                output_path = Path(output_dir) / f"image_{len(images) + 1}.{image_ext}"
                                output_path.parent.mkdir(parents=True, exist_ok=True)
                                with open(output_path, 'wb') as f:
                                    f.write(image_data)
                                image_info["file_path"] = str(output_path)
                            images.append(image_info)

            return images

        except FileNotFoundError as e:
            logger.error(f"File not found: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error extracting images from DOCX {file_path}: {str(e)}")
            raise